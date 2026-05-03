import os
import argparse
import re
import json
import requests
import csv
import yaml
from datetime import datetime
from mmseqs2 import run_mmseqs2
import const
import pandas as pd
import copy
now = datetime.now()
all_line=[]
mem_count=[]
if os.path.exists('neff_data.csv'):
    with open('neff_data.csv','r') as file:
        r=csv.reader(file)
        for line in r:
            all_line.append(line)
class_list={}
check_process=3
# cut=None
cut=4
parts=['first','second']
def _msa_stem(target_id, same_sequence, part_idx=None):
    if same_sequence:
        return target_id
    return f"{target_id}_{parts[part_idx]}"


def _extract_msa_sequences(msa_text):
    if not msa_text:
        return []
    lines = msa_text.strip().splitlines()
    return lines[1::2]


def _build_keyed_rows(paired_text, unpaired_text):
    paired = _extract_msa_sequences(paired_text)[: const.max_paired_seqs]
    keys = [idx for idx, seq in enumerate(paired) if seq != "-" * len(seq)]
    paired = [seq for seq in paired if seq != "-" * len(seq)]

    unpaired = _extract_msa_sequences(unpaired_text)
    unpaired = unpaired[: (const.max_msa_seqs - len(paired))]
    if paired and unpaired:
        unpaired = unpaired[1:]

    seqs = paired + unpaired
    keys.extend([-1] * len(unpaired))
    return keys, seqs


def _write_msa_csv(msa_path, keys, seqs):
    csv_str = ["key,sequence"] + [f"{key},{seq}" for key, seq in zip(keys, seqs)]
    with open(msa_path, "w") as file:
        file.write("\n".join(csv_str))


def _write_a3m(msa_path, query_sequence, seqs):
    a3m_lines = [">query", query_sequence]
    for idx, seq in enumerate(seqs[1:], start=1):
        a3m_lines.extend([f">msa_{idx}", seq])
    with open(msa_path, "w") as file:
        file.write("\n".join(a3m_lines) + "\n")


def compute_msa(
    data,
    target_id: str,
    msa_dir: str,
    msa_server_url: str,
    msa_pairing_strategy: str,
) -> None:
    """Compute the MSA for the input data.

    Parameters
    ----------
    data : dict[str, str]
        The input protein sequences.
    target_id : str
        The target id.
    msa_dir : Path
        The msa directory.
    msa_server_url : str
        The MSA server URL.
    msa_pairing_strategy : str
        The MSA pairing strategy.

    """
    os.makedirs(msa_dir, exist_ok=True)
    same_sequence = len(data) == 1 or all(sequence == data[0] for sequence in data)

    if len(data) > 1:
        paired_msas, paired_chai = run_mmseqs2(
            data,
            os.path.join(msa_dir, f"{target_id}_paired_tmp"),
            use_env=True,
            use_pairing=True,
            host_url=msa_server_url,
            pairing_strategy=msa_pairing_strategy,
        )
    else:
        paired_msas, paired_chai = [""] * len(data), [""] * len(data)

    if not same_sequence:
        for part, sequence in enumerate(data):
            unpaired_msa, unpair_chai = run_mmseqs2(
                sequence,
                os.path.join(msa_dir, f"{target_id}_{parts[part]}_unpaired_tmp"),
                use_env=True,
                use_pairing=False,
                host_url=msa_server_url,
                pairing_strategy=msa_pairing_strategy,
            )
            paired_text = paired_msas[part] if isinstance(paired_msas, list) else paired_msas
            unpaired_text = unpaired_msa[0] if isinstance(unpaired_msa, list) else unpaired_msa
            keys, seqs = _build_keyed_rows(paired_text, unpaired_text)
            stem = _msa_stem(target_id, False, part)
            _write_msa_csv(os.path.join(msa_dir, f"{stem}.csv"), keys, seqs)
            _write_a3m(os.path.join(msa_dir, f"{stem}.a3m"), sequence, seqs)
    else:
        unpaired_msa = run_mmseqs2(
            data[0],
            os.path.join(msa_dir, f"{target_id}_unpaired_tmp"),
            use_env=True,
            use_pairing=False,
            host_url=msa_server_url,
            pairing_strategy=msa_pairing_strategy,
        )
        unpaired_text = unpaired_msa[0] if isinstance(unpaired_msa, list) else unpaired_msa
        paired_text = paired_msas[0] if paired_msas else ""
        keys, seqs = _build_keyed_rows(paired_text, unpaired_text)
        stem = _msa_stem(target_id, True)
        _write_msa_csv(os.path.join(msa_dir, f"{stem}.csv"), keys, seqs)
        _write_a3m(os.path.join(msa_dir, f"{stem}.a3m"), data[0], seqs)
        
def get_protein_sequence(pdb_id,chain,tag):
    sifts_url = f"https://www.ebi.ac.uk/pdbe/api/mappings/uniprot/{pdb_id}"
    response = requests.get(sifts_url).json()
    chain_dic={'A':0,'B':-1}
    try:
        index=chain_dic[chain]
        # if tag==True:
        #     index+=-1
        uniprot_id=list(response[pdb_id.lower()]['UniProt'].keys())[index]
        url = f"https://www.uniprot.org/uniprot/{uniprot_id}.fasta"
        response = requests.get(url)
        sequence=response.text if response.ok else print('manual get')
        current_seq=''.join(re.split(r'SV=\d', sequence)[1].strip('').split('\n'))
    except IndexError:
        sifts_url = f'https://www.rcsb.org/fasta/entry/{pdb_id}/display'
        response = requests.get(sifts_url)
        sequence=response.text.split('Chain ') if response.ok else print('manual get')
        for seq in sequence:
            if seq[0]==chain or 'Chains' in seq:
                current_seq=seq.split('\n')[1]
    return current_seq
def json_process_af3(name, sequences,  multimer='none'):
    json_seq = []
    same_sequence = sequences[0][1] == sequences[1][1]
    
    if multimer == 'memdock':
        # Chain IDs: 4 chains for first sequence, 1 chain for second
        chain_ids = ['A', 'B', 'C', 'D', 'E']
        
        # Create combined FASTA file
        with open(f'new_benchmark/{name}_chai.fasta', 'w') as file:
            for i in range(4):
                file.write(f'>protein|name={name}_{chain_ids[i]}\n{sequences[0][1]}\n')
            file.write(f'>protein|name={name}_{chain_ids[4]}\n{sequences[1][1]}\n')
        
        # Create JSON configuration
        json_seq = [
            {"protein": {"id": chain_ids[:4], "sequence": sequences[0][1],"unpairedMsa":f'msas/{_msa_stem(name, False, 0)}.a3m'}},
            {"protein": {"id": [chain_ids[4]], "sequence": sequences[1][1],"unpairedMsa":f'msas/{_msa_stem(name, False, 1)}.a3m'}}
        ]
        
        # Update base JSON template
        with open('temp.json', 'r') as file:
            data_json = json.load(file)
        data_json['name'] = name
        data_json['sequences'] = json_seq
        data_json['version'] = 2
        with open(f'new_benchmark/{name}.json', 'w') as file:
            json.dump(data_json, file, indent=4)
        
        # Generate YAML configuration
        with open('6ZH1_1.yaml') as f:
            data_yaml = yaml.load(f, Loader=yaml.FullLoader)
        
        new_sequences = []
        # First protein (4 chains)
        for _ in range(4):
            entry = copy.deepcopy(data_yaml['sequences'][0])
            entry['protein']['sequence'] = sequences[0][1]
            entry['protein']['msa'] = f'msas/{name}_{parts[0]}.csv'
            new_sequences.append(entry)
        # Second protein (1 chain)
        entry = copy.deepcopy(data_yaml['sequences'][1])
        entry['protein']['sequence'] = sequences[1][1]
        entry['protein']['msa'] = f'msas/{name}_{parts[1]}.csv'
        new_sequences.append(entry)
        
        data_yaml['sequences'] = new_sequences
        for repeat in range(1000):
            with open(f'new_benchmark/{name}_{repeat+1}.yaml', 'w') as file:
                yaml.dump(data_yaml, file)
    
    elif multimer == 'binder':
        # Chain IDs: 4 chains for each sequence
        chain_ids = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
        
        # Create combined FASTA file
        with open(f'new_benchmark/{name}_chai.fasta', 'w') as file:
            for i in range(4):
                file.write(f'>protein|name={name}_{chain_ids[i]}\n{sequences[0][1]}\n')
            for i in range(4, 8):
                file.write(f'>protein|name={name}_{chain_ids[i]}\n{sequences[1][1]}\n')
        
        # Create JSON configuration
        json_seq = [
            {"protein": {"id": chain_ids[:4], "sequence": sequences[0][1],"unpairedMsa":f'msas/{_msa_stem(name, False, 0)}.a3m'}},
            {"protein": {"id": chain_ids[4:], "sequence": sequences[1][1],"unpairedMsa":f'msas/{_msa_stem(name, False, 1)}.a3m'}}
        ]
        
        # Update base JSON template
        with open('temp.json', 'r') as file:
            data_json = json.load(file)
        data_json['name'] = name
        data_json['sequences'] = json_seq
        data_json['version'] = 2
        with open(f'new_benchmark/{name}.json', 'w') as file:
            json.dump(data_json, file, indent=4)
        
        # Generate YAML configuration
        with open('6ZH1_1.yaml') as f:
            data_yaml = yaml.load(f, Loader=yaml.FullLoader)
        
        new_sequences = []
        # First protein (4 chains)
        for _ in range(4):
            entry = copy.deepcopy(data_yaml['sequences'][0])
            entry['protein']['sequence'] = sequences[0][1]
            entry['protein']['msa'] = f'msas/{name}_{parts[0]}.csv'
            new_sequences.append(entry)
        # Second protein (4 chains)
        for _ in range(4):
            entry = copy.deepcopy(data_yaml['sequences'][1])
            entry['protein']['sequence'] = sequences[1][1]
            entry['protein']['msa'] = f'msas/{name}_{parts[1]}.csv'
            new_sequences.append(entry)
        
        data_yaml['sequences'] = new_sequences
        for repeat in range(1000):
            with open(f'new_benchmark/{name}_{repeat+1}.yaml', 'w') as file:
                yaml.dump(data_yaml, file)
    
    else:
        # Existing dimer handling
        if same_sequence:
            # Homodimer case
            curr_seq = {'protein': {'id': ['A', 'B'], 'sequence': sequences[0][1], 'unpairedMsa': f'msas/{_msa_stem(name, True)}.a3m'}}
            json_seq.append(curr_seq)
            with open(f'new_benchmark/{name}_chai.fasta', 'w') as file:
                file.write(f'>protein|name={name}_A\n{sequences[0][1]}\n')
                file.write(f'>protein|name={name}_B\n{sequences[0][1]}')
        else:
            # Heterodimer case
            for idx, (seq, part) in enumerate(zip(sequences, parts)):
                title = f'{name}_{part}'
                with open(f'new_benchmark/{title}.fasta', 'w') as file:
                    file.write(f'>{title}\n{seq[1]}')
                with open(f'new_benchmark/{name}_chai.fasta', 'a') as file:
                    file.write(f'>protein|name={title}\n{seq[1]}\n')
                curr_seq = {'protein': {'id': [part], 'sequence': seq[1], 'unpairedMsa': f'msas/{_msa_stem(name, False, idx)}.a3m'}}
                json_seq.append(curr_seq)
        
        # Update JSON configuration
        with open('temp.json', 'r') as file:
            data = json.load(file)
        data['name'] = name
        data['sequences'] = json_seq
        data['version'] = 2
        with open(f'new_benchmark/{name}.json', 'w') as file:
            json.dump(data, file, indent=4)
        
        # Generate YAML configurations
        with open('GPR3.yaml') as f:
            data = yaml.load(f, Loader=yaml.FullLoader)
        
        if not same_sequence:
            # Heterodimer YAML
            for index, (seq, part) in enumerate(zip(sequences, parts)):
                data['sequences'][index]['protein']['sequence'] = seq[1]
                data['sequences'][index]['protein']['msa'] = f'msas/{name}_{part}.csv'
            for repeat in range(1000):
                with open(f'new_benchmark/{name}_{repeat+1}.yaml', 'w') as file:
                    yaml.dump(data, file)
        else:
            # Homodimer YAML
            for index, seq in enumerate(sequences):
                data['sequences'][index]['protein']['sequence'] = seq[1]
                data['sequences'][index]['protein']['msa'] = f'msas/{name}.csv'
            for repeat in range(1000):
                with open(f'new_benchmark/{name}_{repeat+1}.yaml', 'w') as file:
                    yaml.dump(data, file)
    msa_seqs=[i[1] for i in sequences]
    compute_msa(
        data=msa_seqs,
        target_id=name,
        msa_dir='msas/',
        msa_server_url="https://api.colabfold.com",
        msa_pairing_strategy='greedy',
    )
    

dest_path="new_benchmark"
month_number_dict = {
    "JAN": '01',
    "FEB": '02',
    "MAR": '03',
    "APR": '04',
    "MAY": '05',
    "JUN": '06',
    "JUL": '07',
    "AUG": '08',
    "SEP": '09',
    "OCT": '10',
    "NOV": '11',
    "DEC": '12'
}
aa3to1={
   'ALA':'A', 'VAL':'V', 'PHE':'F', 'PRO':'P', 'MET':'M',
   'ILE':'I', 'LEU':'L', 'ASP':'D', 'GLU':'E', 'LYS':'K',
   'ARG':'R', 'SER':'S', 'THR':'T', 'TYR':'Y', 'HIS':'H',
   'CYS':'C', 'ASN':'N', 'GLN':'Q', 'TRP':'W', 'GLY':'G',
   'MSE':'M', 'CSO':'C'
}

def fix_missing(seq,miss,pdb_name,chain,a_len):
    sequ_code=''
    # print(seq)
    first=int(seq[0][0])
    # print(first)
    last=int(seq[-1][0])
    block=0
    write_tag=0
    refe_seq=''
    for num,token in enumerate(seq):
        # print(num)
        # print(token)
        count=num+first+block
        if first<=int(token[0])<=last and int(token[0])==count:
            sequ_code+=token[1]
        elif first<=int(token[0])<=last and int(token[0])<count and str(count) in miss.keys():
            lack=miss[str(count)]
            sequ_code+=lack
            sequ_code+=token[1]
            block+=1
            write_tag=1
            
        else:
            write_tag=1
            while int(token[0])!=count:

                if refe_seq=='':
                    tag=False
                    if '-'in list(miss.keys())[0][0]:
                        tag=True
                    refe_seq=get_protein_sequence(pdb_name,chain,tag)
                # print(refe_seq)
                # print(sequ_code)
                lack=refe_seq.split(sequ_code[-5:])[-1][0]
                sequ_code+=lack
                with open(f'new_benchmark/{pdb_name}.patch','a') as patch:
                    patch.write(f'{count} {lack}\n')
                block+=1
                count=num+first+block

            else:
                sequ_code+=token[1]
    if write_tag==1:
        with open(f'new_benchmark/{pdb_name}.patch','a') as patch:
                patch.write(f'chain_A_len:{a_len}\n')
                patch.write(f'chain_{chain}\n')
    return sequ_code
    
def download_pdb(pdb_id, dest_dir):
    url = 'http://www.rcsb.org/pdb/files/%s.pdb' % (pdb_id.upper())
    response = requests.get(url)
    
    if response.status_code == 200:
        print(f'fetching {pdb_id}')
        with open(dest_dir, 'wb') as file:
            file.write(response.content)
    
    else:
        print(f"Request failed with status code: {response.status_code}")
    if "MEMBRANE" in response.text:
        url = f"https://pdbtm.unitmp.org/api/v1/entry/{pdb_id.lower()}.trpdb"
        print(url)
        response = requests.get(url)
        tm_info=response.text if response.ok else 'wrong'
        # print(tm_info)
        if tm_info!='wrong':
            with open(f'new_benchmark/{pdb_id}_m.pdb','w') as file:
                file.write(tm_info)
def read_pdb(pdb_path,pdb_name):
    with open(pdb_path,'r') as file:
        analyze_method=''
        resolution=0
        date=''
        miss_index=0
        molecules={}
        sequences_len=[]
        chain_missing={}
        all_line=file.readlines()
        sequences={'A':{},'B':{}}
        for index,line in enumerate(all_line):
            # if 'COMPND   3 CHAIN:'in line:
            #     chian_list=line.split(':')[-1].split(';')[0].split(',')
            #     # chian_list[0]=chian_list[0].strip(' ')
            #     sequences={chain.strip():{} for chain in chian_list}
            if line.startswith('HEADER'):
                classficiation = line[10:50]
                classficiation = re.sub(r'\s{2,}', ' ', classficiation)
                if classficiation not in class_list.keys():
                    class_list[classficiation]=1
                else:
                    class_list[classficiation]+=1
            elif line.startswith('EXPDTA'):
                if 'SCOPY' in line:
                    analyze_method='Cryo-EM'
                    continue
                elif 'CTION' in line:
                    analyze_method='X-RAY'
                    continue
            elif line.startswith('REVDAT   1'):
                curr=line.split(' ')[6]
                ls=curr.split('-')
                ls[1]=month_number_dict[ls[1]]
                date='-'.join(ls)
                continue
            elif line.startswith('KEYWDS'):
                if 'MEMBRANE' in line and pdb_name not in mem_count:
                    mem_count.append(pdb_name)
            elif line.startswith('REMARK   2 R'):
                resolution = re.split(r'\s+', line)[3]
                continue
            elif line.startswith('FORMUL '):
                curr_mol = re.split(r'\s+', line)
                if '(' in line:

                    molecules[curr_mol[2]]=curr_mol[3].split('(')[0]
                else:
                    molecules[curr_mol[2]]='1'
                # print(molecules)
            elif 'REMARK 465   M RES C SSSEQI' in line:
                miss_index=index
            ca_pattern=re.compile("^(ATOM)\s{0,6}\d{1,5}\s{2}CA\s[\sA]([A-Z]{3})\s([\s\w])\s*(\d{1,4})")
            match_list=ca_pattern.findall(line)
            if match_list:
                resn=match_list[0][1]
                chain=match_list[0][2]
                resi=match_list[0][3]
                # print(match_list)
                # break
                # print(sequences)
                if pdb_name=='7T82' and chain=='F':
                    sequences['B'][resi]=aa3to1[resn]
                elif chain in sequences.keys():
                    sequences[chain][resi]=aa3to1[resn]

        if 'HOH' in molecules.keys():
            del molecules['HOH']
        if  miss_index!=0:
            for line in all_line[miss_index+1:]:
                if 'REMARK 465' in line:
                    curr_miss = re.split(r'\s+', line)
                    if curr_miss[3] not in chain_missing.keys():
                        chain_missing[curr_miss[3]]={curr_miss[4]:aa3to1[curr_miss[2]]}
                    else:
                        chain_missing[curr_miss[3]][curr_miss[4]]=aa3to1[curr_miss[2]]
        # print(chain_missing)
        curr_sequences=[]
        a_len=len(list(sequences['A'].items()))
        
        for chain in sequences.keys():
            whole_list=[aa[1] for aa in sequences[chain].items()]
            sequence=''.join(whole_list)
            if chain in chain_missing.keys():
                sequence=fix_missing(list(sequences[chain].items()),chain_missing[chain],pdb_name,chain,a_len)
                sequences_len.append(str(len(sequence)))
            curr_seq=[chain,sequence]
            curr_sequences.append(curr_seq)
                # print(curr_sequences)
            
        json_process_af3(pdb_name,curr_sequences)
        # print(sequences_len)
    return analyze_method,resolution,date,molecules,sequences_len,classficiation


def process_docx(docx_path, table_index=0, start=0, stop=None):
    from docx import Document
    doc = Document(docx_path)
    table = doc.tables[table_index]
    num_row = len(table.rows)
    end = num_row if stop is None else min(stop, num_row)

    for line in range(start, end):
        pdb_name = table.cell(line, 0).text.strip('\n').strip()
        if not pdb_name:
            continue

        neff = []
        print(f'{pdb_name} running,-------{line}/{num_row-1}-------')
        for neff_line in all_line:
            if pdb_name == neff_line[0]:
                neff = neff_line[1:]
                break

        pdb_path = os.path.join(dest_path, pdb_name + '.pdb')
        if not os.path.isfile(pdb_path):
            download_pdb(pdb_name, pdb_path)

        analyze_method, resolution, date, molecules, sequences_len, classification = read_pdb(pdb_path, pdb_name)
        table.cell(line, 1).text = str(resolution)
        table.cell(line, 2).text = date
        table.cell(line, 3).text = '\n'.join(sequences_len)
        table.cell(line, 4).text = analyze_method
        table.cell(line, 5).text = '\n'.join([f'{molecules[mol]}*{mol}' for mol in molecules.keys()]) if molecules else '-'
        table.cell(line, 6).text = f"class:{classification}\nneff:{neff}\n" + table.cell(line, 6).text

    return doc


def main():
    parser = argparse.ArgumentParser(description='Generate AF3/Chai/Boltz inputs and precompute MSAs from a DOCX table of PDB IDs.')
    parser.add_argument('--docx', help='Path to the input Word document containing the table of PDB IDs.')
    parser.add_argument('--table-index', type=int, default=0, help='Table index in the Word document.')
    parser.add_argument('--start', type=int, default=0, help='Start row index.')
    parser.add_argument('--stop', type=int, help='Stop row index (exclusive).')
    parser.add_argument('--save-docx', help='Optional output Word document path.')
    args = parser.parse_args()

    if not args.docx:
        return

    doc = process_docx(args.docx, args.table_index, args.start, args.stop)
    if args.save_docx:
        doc.save(args.save_docx)

# doc = Document("new structure.docx")
# table = doc.tables[0]

# num_row=len(table.rows)
# print(num_row)
# for line in range(num_row)[check_process:cut]:
#     PDB_name=table.cell(line, 0).text.strip('\n')
#     neff=[]
#     print(f'{PDB_name} running,-------{line}/{num_row-1}-------')
#     for neff_line in all_line:
#         if PDB_name==neff_line[0]:
#             neff=neff_line[1:]
#     if PDB_name!='':
#         pdb_path = os.path.join(dest_path, PDB_name + '.pdb')
#         if not os.path.isfile(pdb_path):
#             download_pdb(PDB_name, pdb_path)
#         analyze_method, resolution, date, molecules, sequences_len, classification = read_pdb(pdb_path, PDB_name)
#         table.cell(line, 1).text = str(resolution)
#         table.cell(line, 2).text = date
#         table.cell(line, 3).text = '\n'.join(sequences_len)
#         table.cell(line, 4).text = analyze_method
#         table.cell(line, 5).text = '\n'.join([f'{molecules[mol]}*{mol}' for mol in molecules.keys()]) if molecules else '-'
#         table.cell(line, 6).text = f"class:{classification}\nneff:{neff}\n"+table.cell(line, 6).text 

#  os.system(f'tar -jcvf {now.strftime("%y_%m_%d")}_boltz.tar.bz2 new_benchmark/*.yaml')

# os.system(f'tar -jcvf {now.strftime("%y_%m_%d")}_msas.tar.bz2 msas/ --exclude="*env/"')
# doc.save("C:\\project Complex\\new structure.docx")
# print(class_list)

# os.system(f'tar -jcvf {now.strftime("%y_%m_%d")}_rosetta.tar.bz2 .\\new_benchmark\\*_?.fasta .\\new_benchmark\\*.patch .\\new_benchmark\\*.pdb')
# os.system(f'tar -jcvf {now.strftime("%y_%m_%d")}_rosetta.tar.bz2 .\\new_benchmark\\*.patch')
# os.system(f'tar -jcvf {now.strftime("%y_%m_%d")}_chai.tar.bz2 .\\new_benchmark\\*chai.fasta .\\new_benchmark\\*.json')
# print(mem_count)
# print(len(mem_count))


if __name__ == '__main__':
    main()
